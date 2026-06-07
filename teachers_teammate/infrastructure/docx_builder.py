"""
DOCX document creator: implements :class:`~teachers_teammate.interfaces.DocumentCreator`.

Produces a landscape Word document in one of two layouts:

* ``table``    — three columns: Original Image | OCR Result | Correction Proposal
* ``comments`` — two columns: Original Image | OCR Result, with Word diff-comments
"""

from __future__ import annotations

from datetime import UTC, datetime
import difflib
from itertools import zip_longest
import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

from ..interfaces import DocumentCreator

_logger = logging.getLogger(__name__)

_HEADER_FILL = "D9D9D9"  # light grey
_HEADING_COLOUR = RGBColor(0x1F, 0x49, 0x7D)  # Word dark-blue
_IMG_COL_WIDTH = Inches(2.8)  # original-image column
_TXT_COL_WIDTH = Inches(3.2)  # each text column
# Comments format (2-column)
_COMMENTS_WIDE_COL = _IMG_COL_WIDTH + _TXT_COL_WIDTH
_COMMENTS_URI = PackURI("/word/comments.xml")
_COMMENTS_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
_COMMENTS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_COMMENTS_AUTHOR = "Teacher's Teammate"


class DocxDocumentCreator(DocumentCreator):
    """Builds a landscape Word document with an image + OCR + correction table."""

    def __init__(self, *, fmt: str = "table") -> None:
        self._fmt = fmt

    def create(
        self,
        raw_text: str,
        corrected_text: str | None,
        out_path: Path,
        title: str,
        *,
        source_image: str | Path | None = None,
    ) -> None:
        """Write a landscape .docx to *out_path*.

        Dispatches to the table or comments layout based on the ``fmt`` passed
        to the constructor.

        Args:
            raw_text:       Raw OCR output.
            corrected_text: Proofread text, or ``None`` when correction was skipped.
            out_path:       Destination path (parent directory must exist).
            title:          Document heading.
            source_image:   Path to the source image shown in column 1.
        """
        if self._fmt == "comments":
            self._build_with_comments(
                raw_text, corrected_text, out_path, title, image_path=source_image
            )
        else:
            self._build_table(raw_text, corrected_text, out_path, title, image_path=source_image)

    def _build_table(
        self,
        raw_text: str,
        corrected_text: str | None,
        out_path: Path,
        title: str,
        *,
        image_path: str | Path | None = None,
    ) -> None:
        """
        Write a landscape three-column .docx file to *out_path*.

        Args:
            raw_text:       Raw OCR output.
            corrected_text: Proofread text, or None when correction was skipped.
            out_path:       Destination path (parent directory must exist).
            title:          Document heading (typically the source file stem).
            image_path:     Path to the original source image shown in column 1.
                            When ``None``, the image cell is left blank.
        """
        doc = Document()

        # Landscape orientation
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        heading = doc.add_heading(title, level=1)
        for run in heading.runs:
            run.font.color.rgb = _HEADING_COLOUR

        table = doc.add_table(rows=2, cols=3)
        table.style = "Table Grid"
        table.columns[0].width = _IMG_COL_WIDTH
        table.columns[1].width = _TXT_COL_WIDTH
        table.columns[2].width = _TXT_COL_WIDTH

        hdr = table.rows[0].cells
        hdr[0].width = _IMG_COL_WIDTH
        hdr[1].width = _TXT_COL_WIDTH
        hdr[2].width = _TXT_COL_WIDTH
        self._fill_cell(hdr[0], "Original", bold=True)
        self._fill_cell(hdr[1], "OCR Result", bold=True)
        self._fill_cell(hdr[2], "Correction Proposal", bold=True)
        for cell in hdr:
            self._shade_cell(cell, _HEADER_FILL)

        content = table.rows[1].cells
        content[0].width = _IMG_COL_WIDTH
        content[1].width = _TXT_COL_WIDTH
        content[2].width = _TXT_COL_WIDTH
        if image_path is not None:
            self._insert_image(content[0], str(image_path), _IMG_COL_WIDTH)
        self._fill_cell(content[1], raw_text)
        placeholder = corrected_text if corrected_text else "(correction skipped)"
        self._fill_cell(content[2], placeholder, italic=not corrected_text)

        doc.save(str(out_path))

    def _build_with_comments(
        self,
        raw_text: str,
        corrected_text: str | None,
        out_path: Path,
        title: str,
        *,
        image_path: str | Path | None = None,
    ) -> None:
        """
        Write a landscape .docx to *out_path* (two-column: image + OCR text).

        When *corrected_text* is provided, differences against the raw OCR are
        highlighted in yellow and attached as Word comments so reviewers can
        accept or reject each change directly in Word.

        Args:
            raw_text:       Raw OCR output.
            corrected_text: Proofread text used to compute diffs, or None.
            out_path:       Destination path (parent directory must exist).
            title:          Document heading.
            image_path:     Source image embedded in column 1.
        """
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        heading = doc.add_heading(title, level=1)
        for run in heading.runs:
            run.font.color.rgb = _HEADING_COLOUR

        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.columns[0].width = _IMG_COL_WIDTH
        table.columns[1].width = _COMMENTS_WIDE_COL

        hdr = table.rows[0].cells
        hdr[0].width = _IMG_COL_WIDTH
        hdr[1].width = _COMMENTS_WIDE_COL
        self._fill_cell(hdr[0], "Original", bold=True)
        self._fill_cell(hdr[1], "OCR Result (corrections as comments)", bold=True)
        for cell in hdr:
            self._shade_cell(cell, _HEADER_FILL)

        content = table.rows[1].cells
        content[0].width = _IMG_COL_WIDTH
        content[1].width = _COMMENTS_WIDE_COL
        if image_path is not None:
            self._insert_image(content[0], str(image_path), _IMG_COL_WIDTH)

        pending: list[tuple[int, str]] = []
        if corrected_text:
            self._fill_cell_with_diffs(content[1], raw_text, corrected_text, pending)
        else:
            self._fill_cell(content[1], raw_text)
        if pending:
            self._attach_comments_part(doc, pending)

        doc.save(str(out_path))

    @staticmethod
    def _shade_cell(cell: Any, fill_hex: str) -> None:
        tc = cell._tc  # noqa: SLF001  # python-docx has no public API for cell shading; _tc is the only access path to the underlying lxml element
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tc_pr.append(shd)

    @staticmethod
    def _insert_image(cell: Any, image_path: str, max_width: Any) -> None:
        """Insert *image_path* into *cell* scaled to *max_width*."""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        try:
            run.add_picture(image_path, width=max_width)
        except Exception as exc:  # noqa: BLE001  # PIL / python-docx image errors vary by format; log the warning and leave placeholder text
            _logger.warning("Could not insert image '%s': %s", image_path, exc)
            run.text = "(image unavailable)"

    @staticmethod
    def _fill_cell(
        cell: Any,
        text: str,
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(10)

    def _fill_cell_with_diffs(
        self,
        cell: Any,
        raw: str,
        corrected: str,
        comments: list[tuple[int, str]],
    ) -> None:
        """Fill *cell* with *raw*, annotating diffs against *corrected* as Word comments."""
        raw_lines = raw.splitlines()
        corr_lines = corrected.splitlines()
        matcher = difflib.SequenceMatcher(None, raw_lines, corr_lines, autojunk=False)
        first = True
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            raw_block = raw_lines[i1:i2]
            corr_block = corr_lines[j1:j2]
            if tag == "equal":
                for line in raw_block:
                    para = cell.paragraphs[0] if first else cell.add_paragraph()
                    first = False
                    self._add_plain_run(para, line)
            else:
                for raw_line, corr_line in zip_longest(raw_block, corr_block, fillvalue=""):
                    para = cell.paragraphs[0] if first else cell.add_paragraph()
                    first = False
                    self._add_line_with_word_diff(para, raw_line, corr_line, comments)

    def _add_line_with_word_diff(
        self,
        para: Any,
        raw_line: str,
        corr_line: str,
        comments: list[tuple[int, str]],
    ) -> None:
        """Add *raw_line* to *para*, marking word-level diffs against *corr_line* as comments."""
        raw_words = raw_line.split()
        corr_words = corr_line.split()
        if not raw_words and not corr_words:
            return
        matcher = difflib.SequenceMatcher(None, raw_words, corr_words, autojunk=False)
        parts: list[tuple[str, str | None]] = []
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            raw_chunk = " ".join(raw_words[i1:i2])
            corr_chunk = " ".join(corr_words[j1:j2])
            if tag == "equal":
                parts.append((raw_chunk, None))
            elif tag == "replace":
                parts.append((raw_chunk, f'\u2192 "{corr_chunk}"'))
            elif tag == "delete":
                parts.append((raw_chunk, "\u2192 (delete)"))
            else:  # insert — no raw anchor; annotate as inline marker with a comment
                parts.append((f"[+{corr_chunk}]", f'\u2192 (insert "{corr_chunk}")'))
        for idx, (text, comment_text) in enumerate(parts):
            spaced = text + (" " if idx < len(parts) - 1 else "")
            if comment_text is None:
                self._add_plain_run(para, spaced)
            else:
                cid = len(comments)
                comments.append((cid, comment_text))
                self._add_commented_run(para, spaced, cid)

    @staticmethod
    def _add_plain_run(para: Any, text: str) -> None:
        run = para.add_run(text)
        run.font.size = Pt(10)

    @staticmethod
    def _add_commented_run(para: Any, text: str, comment_id: int) -> None:
        """Add a yellow-highlighted run in *para* anchored to *comment_id*."""
        range_start = OxmlElement("w:commentRangeStart")
        range_start.set(qn("w:id"), str(comment_id))
        run = para.add_run(text)
        run.font.size = Pt(10)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run._r.addprevious(range_start)  # noqa: SLF001  # python-docx has no public API for OOXML comment range anchors; _r exposes the underlying lxml element
        range_end = OxmlElement("w:commentRangeEnd")
        range_end.set(qn("w:id"), str(comment_id))
        run._r.addnext(range_end)  # noqa: SLF001  # python-docx has no public API for OOXML comment range anchors; _r exposes the underlying lxml element
        ref_r = OxmlElement("w:r")
        ref_rpr = OxmlElement("w:rPr")
        ref_style = OxmlElement("w:rStyle")
        ref_style.set(qn("w:val"), "CommentReference")
        ref_rpr.append(ref_style)
        ref_r.append(ref_rpr)
        ref_ref = OxmlElement("w:commentReference")
        ref_ref.set(qn("w:id"), str(comment_id))
        ref_r.append(ref_ref)
        range_end.addnext(ref_r)

    @staticmethod
    def _attach_comments_part(doc: Any, comments: list[tuple[int, str]]) -> None:
        """Build word/comments.xml and attach it to *doc* as a package part."""
        root = etree.Element(f"{{{_W_NS}}}comments", nsmap={"w": _W_NS})
        for cid, ctext in comments:
            comment_el = etree.SubElement(root, f"{{{_W_NS}}}comment")
            comment_el.set(f"{{{_W_NS}}}id", str(cid))
            comment_el.set(f"{{{_W_NS}}}author", _COMMENTS_AUTHOR)
            comment_el.set(f"{{{_W_NS}}}date", datetime.now(UTC).strftime("%Y-%m-%dT%H:%M:%SZ"))
            comment_el.set(f"{{{_W_NS}}}initials", "OCR")
            para_el = etree.SubElement(comment_el, f"{{{_W_NS}}}p")
            run_el = etree.SubElement(para_el, f"{{{_W_NS}}}r")
            t_el = etree.SubElement(run_el, f"{{{_W_NS}}}t")
            t_el.text = ctext
            t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        xml_bytes = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
        part = Part(_COMMENTS_URI, _COMMENTS_CT, xml_bytes, doc.part.package)
        doc.part.relate_to(part, _COMMENTS_REL)
