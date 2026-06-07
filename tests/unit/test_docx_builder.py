"""Unit tests for teachers_teammate.docx_builder."""
# pylint: disable=W0404,W0621,W0613  # reimported — monkeypatch/patch blocks locally reimport the patched symbol / redefined-outer-name — pytest fixtures shadow module-scope names by design / unused-argument — pytest injects fixtures by parameter name; not all are used in every test

from __future__ import annotations

from pathlib import Path

from docx import Document
import pytest

from teachers_teammate.infrastructure.docx_builder import DocxDocumentCreator

# ── table format ───────────────────────────────────────────────────────────


@pytest.mark.use_case("DOCX_Report_Export")
def test_create_table_format_produces_docx_file(tmp_path: Path, sample_png: Path) -> None:
    """
    Given  a DocxDocumentCreator in table format with raw + corrected text
    When   create() is called
    Then   a .docx file exists at out_path
    """
    out = tmp_path / "test.docx"
    creator = DocxDocumentCreator(fmt="table")
    creator.create(
        raw_text="raw ocr output",
        corrected_text="corrected output",
        out_path=out,
        title="Test Page",
        source_image=sample_png,
    )
    assert out.is_file()
    assert out.stat().st_size > 0


def test_create_table_format_without_image(tmp_path: Path) -> None:
    """
    Given  no source_image supplied
    When   create() is called in table format
    Then   a valid .docx file is produced (image cell left blank)
    """
    out = tmp_path / "no_image.docx"
    creator = DocxDocumentCreator(fmt="table")
    creator.create(
        raw_text="text",
        corrected_text=None,
        out_path=out,
        title="No Image",
    )
    assert out.is_file()


def test_create_table_format_contains_title(tmp_path: Path) -> None:
    """
    Given  title='My Title'
    When   create() produces a table-format DOCX
    Then   'My Title' appears in the document's paragraphs
    """
    out = tmp_path / "titled.docx"
    creator = DocxDocumentCreator(fmt="table")
    creator.create(
        raw_text="some text",
        corrected_text="corrected",
        out_path=out,
        title="My Title",
    )
    doc = Document(str(out))
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "My Title" in all_text


def test_create_table_format_contains_raw_and_corrected_text(tmp_path: Path) -> None:
    """
    Given  raw_text='raw content' and corrected_text='corrected content'
    When   create() produces a table-format DOCX
    Then   both strings appear somewhere in the document
    """
    out = tmp_path / "content.docx"
    creator = DocxDocumentCreator(fmt="table")
    creator.create(
        raw_text="raw content",
        corrected_text="corrected content",
        out_path=out,
        title="Test",
    )
    doc = Document(str(out))
    all_text = " ".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert "raw content" in all_text
    assert "corrected content" in all_text


def test_create_table_with_none_corrected_text(tmp_path: Path) -> None:
    """
    Given  corrected_text=None (correction disabled)
    When   create() is called
    Then   a .docx file is produced without error
    """
    out = tmp_path / "no_corr.docx"
    creator = DocxDocumentCreator(fmt="table")
    creator.create(raw_text="raw", corrected_text=None, out_path=out, title="T")
    assert out.is_file()


# ── comments format ────────────────────────────────────────────────────────


@pytest.mark.use_case("DOCX_Report_Export")
def test_create_comments_format_produces_docx_file(tmp_path: Path) -> None:
    """
    Given  a DocxDocumentCreator in comments format
    When   create() is called with identical raw and corrected text
    Then   a .docx file is produced
    """
    out = tmp_path / "comments.docx"
    creator = DocxDocumentCreator(fmt="comments")
    creator.create(
        raw_text="hello world",
        corrected_text="hello world",
        out_path=out,
        title="Comments Test",
    )
    assert out.is_file()
    assert out.stat().st_size > 0


def test_create_comments_format_with_diff(tmp_path: Path) -> None:
    """
    Given  raw_text and corrected_text that differ
    When   create() is called in comments format
    Then   a .docx is produced without error
    """
    out = tmp_path / "diff.docx"
    creator = DocxDocumentCreator(fmt="comments")
    creator.create(
        raw_text="The qick brown fox",
        corrected_text="The quick brown fox",
        out_path=out,
        title="Diff Test",
    )
    assert out.is_file()


def test_create_comments_format_with_multiline_diff_attaches_comments_xml(
    tmp_path: Path,
) -> None:
    """
    Given  multi-line raw and corrected text with many differences
    When   create() is called in comments format
    Then   a .docx is produced with a comments.xml part attached
    """
    import zipfile  # noqa: PLC0415

    out = tmp_path / "multi_diff.docx"
    creator = DocxDocumentCreator(fmt="comments")
    creator.create(
        raw_text="helo world\nthis iz rong\nextra deleted line",
        corrected_text="hello world\nthis is correct\nnew inserted line",
        out_path=out,
        title="MultiDiff",
    )
    assert out.is_file()
    with zipfile.ZipFile(out) as z:
        names = z.namelist()
    assert any("comment" in n.lower() for n in names)


def test_create_comments_format_with_image(tmp_path: Path, sample_png: Path) -> None:
    """
    Given  fmt='comments', differing raw and corrected text, and a source image
    When   create() is called
    Then   a .docx is written with the image embedding code path exercised
    """
    out = tmp_path / "img_comments.docx"
    creator = DocxDocumentCreator(fmt="comments")
    creator.create(
        raw_text="hello wrold",
        corrected_text="hello world",
        out_path=out,
        title="Doc",
        source_image=sample_png,
    )
    assert out.is_file()


def test_insert_image_writes_unavailable_text_on_failure(tmp_path: Path) -> None:
    """
    Given  an image path that does not exist (triggers add_picture exception)
    When   _insert_image() is called directly
    Then   no exception propagates and the run text is set to '(image unavailable)'
    """
    from docx import Document  # noqa: PLC0415
    from docx.shared import Inches  # noqa: PLC0415

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    DocxDocumentCreator._insert_image(cell, "/nonexistent/path.png", Inches(2))
    run_texts = [r.text for r in cell.paragraphs[0].runs]
    assert any("unavailable" in t for t in run_texts)


# ── Internal helpers ───────────────────────────────────────────────────────


def test_fill_cell_with_diffs_equal_lines_no_comments(tmp_path: Path) -> None:
    """
    Given  identical raw and corrected multi-line text
    When   _fill_cell_with_diffs() is called
    Then   the pending comments list remains empty
    """
    from docx import Document  # noqa: PLC0415

    creator = DocxDocumentCreator(fmt="comments")
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    text = "line one\nline two\nline three"
    pending: list = []
    creator._fill_cell_with_diffs(cell, text, text, pending)
    assert pending == []


def test_fill_cell_with_diffs_replace_adds_comments(tmp_path: Path) -> None:
    """
    Given  raw and corrected text with a replaced line
    When   _fill_cell_with_diffs() is called
    Then   at least one comment is appended to the pending list
    """
    from docx import Document  # noqa: PLC0415

    creator = DocxDocumentCreator(fmt="comments")
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    pending: list = []
    creator._fill_cell_with_diffs(cell, "old line\nend", "new line\nend", pending)
    assert len(pending) > 0


def test_add_line_with_word_diff_handles_insert_opcode(tmp_path: Path) -> None:
    """
    Given  a raw line with fewer words than the corrected line (insert opcode)
    When   _add_line_with_word_diff() is called
    Then   a comment is added for the inserted words
    """
    from docx import Document  # noqa: PLC0415

    creator = DocxDocumentCreator(fmt="comments")
    doc = Document()
    para = doc.add_paragraph()
    comments: list = []
    creator._add_line_with_word_diff(para, "hello", "hello world here", comments)
    assert len(comments) > 0


def test_add_line_with_word_diff_handles_delete_opcode(tmp_path: Path) -> None:
    """
    Given  a raw line with more words than the corrected line (delete opcode)
    When   _add_line_with_word_diff() is called
    Then   a comment with '→ (delete)' text is added
    """
    from docx import Document  # noqa: PLC0415

    creator = DocxDocumentCreator(fmt="comments")
    doc = Document()
    para = doc.add_paragraph()
    comments: list = []
    creator._add_line_with_word_diff(para, "hello extra words", "hello", comments)
    assert any("delete" in c[1].lower() for c in comments)


def test_add_line_with_word_diff_both_empty_is_noop(tmp_path: Path) -> None:
    """
    Given  both raw_line and corr_line are empty strings
    When   _add_line_with_word_diff() is called
    Then   no exception is raised and no comments are added
    """
    from docx import Document  # noqa: PLC0415

    creator = DocxDocumentCreator(fmt="comments")
    doc = Document()
    para = doc.add_paragraph()
    comments: list = []
    creator._add_line_with_word_diff(para, "", "", comments)
    assert comments == []


def test_attach_comments_part_adds_relationship_to_document(tmp_path: Path) -> None:
    """
    Given  a Document and a list of (comment_id, text) pairs
    When   _attach_comments_part() is called
    Then   the document part has a relationship pointing to comments XML
    """
    from docx import Document  # noqa: PLC0415

    doc = Document()
    DocxDocumentCreator._attach_comments_part(doc, [(0, "→ correction"), (1, "→ (delete)")])
    rels = doc.part.rels
    assert any("comment" in str(v.reltype).lower() for v in rels.values())


def test_create_comments_format_deleted_lines(tmp_path: Path) -> None:
    """
    Given  raw text has extra lines not in corrected (delete opcode at line level)
    When   create() in comments format
    Then   a .docx is produced without error
    """
    out = tmp_path / "deleted.docx"
    creator = DocxDocumentCreator(fmt="comments")
    creator.create(
        raw_text="keep this\ndelete this\nkeep also",
        corrected_text="keep this\nkeep also",
        out_path=out,
        title="Deleted",
    )
    assert out.is_file()


def test_create_comments_format_inserted_lines(tmp_path: Path) -> None:
    """
    Given  corrected text has extra lines not in raw (insert opcode at line level)
    When   create() in comments format
    Then   a .docx is produced without error
    """
    out = tmp_path / "inserted.docx"
    creator = DocxDocumentCreator(fmt="comments")
    creator.create(
        raw_text="first line\nlast line",
        corrected_text="first line\nnew middle line\nlast line",
        out_path=out,
        title="Inserted",
    )
    assert out.is_file()


# ── private-API smoke tests (B5 pin guard) ─────────────────────────────────
# These tests verify that the python-docx private attributes used in docx_builder
# (_tc for cell shading, _r for comment range anchors) still work after package
# updates.  If a python-docx upgrade renames or removes these attributes, the
# tests fail here — before any user is affected.


def test_table_format_round_trip_is_parseable_with_shaded_header(tmp_path: Path) -> None:
    """Given raw and corrected text, when table format is written, then the .docx is
    readable by python-docx and contains a table with at least two header cells."""
    out = tmp_path / "smoke_table.docx"
    DocxDocumentCreator(fmt="table").create("raw text", "corrected text", out, "title")

    doc = Document(str(out))
    assert len(doc.tables) >= 1
    assert len(doc.tables[0].rows[0].cells) >= 2


def test_comments_format_round_trip_is_parseable(tmp_path: Path) -> None:
    """Given raw and corrected text, when comments format is written, then the .docx is
    readable by python-docx and contains paragraphs."""
    out = tmp_path / "smoke_comments.docx"
    DocxDocumentCreator(fmt="comments").create("raw text", "corrected text", out, "title")

    assert Document(str(out)).paragraphs
